#!/usr/bin/env python3
# filepath: tools/ocr_openai.py
"""
Image/PDF → Text via OpenAI Vision (Responses API) with GUI + Batch + Multiple export formats.

End‑user friendly:
  • Double‑click opens a GUI (Tkinter). No extra GUI deps.
  • Batch mode: add many files and/or a folder; outputs go to a chosen directory.
  • Export formats: .txt, .md, .docx

CLI examples:
  python tools/ocr_openai.py --gui
  python tools/ocr_openai.py input.pdf -o out.txt --format txt
  python tools/ocr_openai.py img1.jpg img2.png --format docx --out-dir ./out
  python tools/ocr_openai.py --folder ./scans --format md --out-dir ./out

Install:␊
    pip install openai==1.* pymupdf python-docx
    setx OPENAI_API_KEY your_key   # Windows
    export OPENAI_API_KEY=your_key # macOS/Linux

Notes:
  - Rendering PDFs with PyMuPDF avoids external system tools.
  - Verbatim prompt minimizes paraphrasing by the model.
"""
from __future__ import annotations

import argparse
import base64
import mimetypes
import os
import queue
import sys
import threading
from dataclasses import dataclass
from pathlib import Path
from typing import Iterable, List, Optional, Sequence, Tuple

# Third‑party runtime deps
try:
    import fitz  # PyMuPDF
except ImportError:  # pragma: no cover
    fitz = None  # type: ignore

try:
    from openai import OpenAI
except ImportError:  # pragma: no cover
    OpenAI = None  # type: ignore

# Optional for DOCX export; checked at runtime only if used
try:
    from docx import Document  # type: ignore
except ImportError:
    Document = None  # type: ignore

# Stdlib GUI
import tkinter as tk
from tkinter import ttk, filedialog, messagebox


# ---------------------------- Config & Prompt ----------------------------
DEFAULT_MODEL = os.environ.get("OCR_OPENAI_MODEL", "gpt-4o-mini")
DEFAULT_DPI = 220
SUPPORTED_IMAGE_EXTS = {".png", ".jpg", ".jpeg", ".bmp", ".tif", ".tiff", ".webp"}
SUPPORTED_INPUT_EXTS = SUPPORTED_IMAGE_EXTS | {".pdf"}
SUPPORTED_FORMATS = ("txt", "md", "docx")
# Safer newline constant to avoid accidental line breaks in concatenations
NL = chr(10)  # newline character

VERBATIM_PROMPT = (
    "Extract the text verbatim from this page."
    " Preserve original line breaks, spacing, punctuation, and casing."
    " Do NOT summarize, reflow, interpret, or add words."
    " If characters are unreadable, output '�' for those characters."
    " Output plain UTF-8 text only with no explanations."
)


# ---------------------------- Data Types ----------------------------
@dataclass
class ImageBytes:
    data: bytes
    mime: str
    page_index: Optional[int] = None


# ---------------------------- I/O Helpers ----------------------------
def guess_mime(path: Path) -> str:
    mime, _ = mimetypes.guess_type(str(path))
    if mime is None:
        return "image/png" if path.suffix.lower() == ".png" else "image/jpeg"
    return mime


def load_image_file(path: Path) -> ImageBytes:
    mime = guess_mime(path)
    with path.open("rb") as f:
        data = f.read()
    return ImageBytes(data=data, mime=mime)


def pdf_page_count(path: Path) -> int:
    if fitz is None:
        raise RuntimeError("PyMuPDF (fitz) is required. Install with: pip install pymupdf")
    doc = fitz.open(str(path))
    n = doc.page_count
    doc.close()
    return n


def pdf_pages_as_images(path: Path, dpi: int = DEFAULT_DPI) -> Iterable[ImageBytes]:
    if fitz is None:
        raise RuntimeError("PyMuPDF (fitz) is required. Install with: pip install pymupdf")
    doc = fitz.open(str(path))
    zoom = dpi / 72.0
    mat = fitz.Matrix(zoom, zoom)
    try:
        for i, page in enumerate(doc):
            pix = page.get_pixmap(matrix=mat, alpha=False)
            yield ImageBytes(data=pix.tobytes(output="png"), mime="image/png", page_index=i)
    finally:
        doc.close()


def encode_image_data_url(img: ImageBytes) -> str:
    b64 = base64.b64encode(img.data).decode("utf-8")
    return f"data:{img.mime};base64,{b64}"


# ---------------------------- OpenAI Call ----------------------------
class OpenAIOCR:
    def __init__(self, model: str = DEFAULT_MODEL, api_key: Optional[str] = None):
        if OpenAI is None:  # pragma: no cover - handled via runtime check
            raise RuntimeError("openai SDK is required. Install with: pip install openai")
        self.client = OpenAI(api_key=api_key) if api_key else OpenAI()
        self.model = model

    def ocr_image(self, img: ImageBytes, prompt: str = VERBATIM_PROMPT) -> str:
        data_url = encode_image_data_url(img)
        response = self.client.responses.create(
            model=self.model,
            input=[
                {
                    "role": "user",
                    "content": [
                        {"type": "input_text", "text": prompt},
                        {"type": "input_image", "image_url": data_url},
                    ],
                }
            ],
        )
        return getattr(response, "output_text", "").rstrip()


# ---------------------------- Conversion/Export ----------------------------

def page_header(page_num: int) -> str:
    """Return the standard page header used in the plain-text output."""
    return f"{NL}{NL}===== Page {page_num} ====={NL}"


def normalize_to_markdown(text: str) -> str:
    lines: List[str] = []
    for line in text.splitlines():
        if line.startswith("===== Page ") and line.rstrip().endswith("====="):
            title = line.strip(" =")
            lines.append(f"## {title}")
        else:
            lines.append(line)
    return "\n".join(lines).rstrip() + "\n"


def write_output(text: str, out_path: Path, fmt: str) -> None:
    fmt = fmt.lower()
    out_path.parent.mkdir(parents=True, exist_ok=True)
    if fmt == "txt":
        out_path.write_text(text, encoding="utf-8")
        return
    if fmt == "md":
        out_path.write_text(normalize_to_markdown(text), encoding="utf-8")
        return
    if fmt == "docx":
        if Document is None:
            raise RuntimeError("python-docx not installed. Run: pip install python-docx")
        doc = Document()
        # Minimal styling; keep verbatim feel
        for line in text.splitlines():
            if line.startswith("===== Page ") and line.rstrip().endswith("====="):
                doc.add_heading(line.strip(" ="), level=2)
            elif line.strip() == "":
                doc.add_paragraph("")
            else:
                doc.add_paragraph(line)
        doc.save(str(out_path))
        return
    raise ValueError(f"Unsupported format: {fmt}")


def target_path_for(inp: Path, out_dir: Path, fmt: str) -> Path:
    ext = "." + fmt.lower()
    return out_dir.joinpath(inp.with_suffix(ext).name)


# ---------------------------- Core Orchestration ----------------------------

def ocr_single_file(
    inp: Path,
    ocr: OpenAIOCR,
    dpi: int,
    on_page: Optional[callable] = None,
    cancel: Optional[threading.Event] = None,
) -> str:
    if inp.suffix.lower() == ".pdf":
        total = pdf_page_count(inp)
        texts: List[str] = []
        for img in pdf_pages_as_images(inp, dpi=dpi):
            if cancel and cancel.is_set():
                raise RuntimeError("Operation cancelled")
            if on_page:
                on_page((img.page_index or 0) + 1, total)
            header = page_header((img.page_index or 0) + 1)
            page_text = ocr.ocr_image(img)
            texts.append(header + page_text)
        return "".join(texts).lstrip()
    img = load_image_file(inp)
    if on_page:
        on_page(1, 1)
    return ocr.ocr_image(img)


# ---------------------------- GUI ----------------------------
class OCRApp(tk.Tk):
    """GUI for single and batch OCR with format export."""

    def __init__(self) -> None:
        super().__init__()
        self.title("Image/PDF → Text (OpenAI)")
        self.geometry("880x640")
        self.resizable(True, True)

        # State
        self.cancel_event = threading.Event()
        self.worker: Optional[threading.Thread] = None
        self.log_q: "queue.Queue[str]" = queue.Queue()

        # Vars
        self.var_batch = tk.BooleanVar(value=False)
        self.var_input = tk.StringVar()
        self.var_output = tk.StringVar()
        self.var_output_dir = tk.StringVar()
        self.var_format = tk.StringVar(value="txt")
        self.var_model = tk.StringVar(value=DEFAULT_MODEL)
        self.var_dpi = tk.IntVar(value=DEFAULT_DPI)
        self.var_open_when_done = tk.BooleanVar(value=True)
        self.var_api_key = tk.StringVar(value=os.environ.get("OPENAI_API_KEY", ""))

        # Layout
        self._build_ui()
        self._poll_log()

    # UI
    def _build_ui(self) -> None:
        pad = {"padx": 10, "pady": 6}
        frm = ttk.Frame(self)
        frm.pack(fill=tk.BOTH, expand=True)

        # Mode toggle
        row0 = ttk.Frame(frm)
        row0.pack(fill=tk.X, **pad)
        ttk.Checkbutton(row0, text="Batch mode", variable=self.var_batch, command=self._toggle_mode).pack(anchor=tk.W)

        # Single-file section
        self.grp_single = ttk.LabelFrame(frm, text="Single file")
        self.grp_single.pack(fill=tk.X, **pad)
        r1 = ttk.Frame(self.grp_single)
        r1.pack(fill=tk.X, **pad)
        ttk.Label(r1, text="Input file (PDF or image)").pack(anchor=tk.W)
        r1b = ttk.Frame(r1)
        r1b.pack(fill=tk.X)
        ttk.Entry(r1b, textvariable=self.var_input).pack(side=tk.LEFT, fill=tk.X, expand=True)
        ttk.Button(r1b, text="Browse…", command=self._choose_input).pack(side=tk.LEFT, padx=6)

        r2 = ttk.Frame(self.grp_single)
        r2.pack(fill=tk.X, **pad)
        ttk.Label(r2, text="Output file").pack(anchor=tk.W)
        r2b = ttk.Frame(r2)
        r2b.pack(fill=tk.X)
        ttk.Entry(r2b, textvariable=self.var_output).pack(side=tk.LEFT, fill=tk.X, expand=True)
        ttk.Button(r2b, text="Save As…", command=self._choose_output).pack(side=tk.LEFT, padx=6)

        # Batch section
        self.grp_batch = ttk.LabelFrame(frm, text="Batch queue")
        self.grp_batch.pack(fill=tk.BOTH, expand=True, **pad)
        r3t = ttk.Frame(self.grp_batch)
        r3t.pack(fill=tk.X, **pad)
        ttk.Button(r3t, text="Add Files…", command=self._add_files).pack(side=tk.LEFT)
        ttk.Button(r3t, text="Add Folder…", command=self._add_folder).pack(side=tk.LEFT, padx=(8, 0))
        ttk.Button(r3t, text="Remove Selected", command=self._remove_selected).pack(side=tk.LEFT, padx=(8, 0))
        ttk.Button(r3t, text="Clear", command=self._clear_list).pack(side=tk.LEFT, padx=(8, 0))

        r3 = ttk.Frame(self.grp_batch)
        r3.pack(fill=tk.BOTH, expand=True, **pad)
        self.lst_inputs = tk.Listbox(r3, selectmode=tk.EXTENDED)
        self.lst_inputs.pack(side=tk.LEFT, fill=tk.BOTH, expand=True)
        sb = ttk.Scrollbar(r3, orient=tk.VERTICAL, command=self.lst_inputs.yview)
        sb.pack(side=tk.RIGHT, fill=tk.Y)
        self.lst_inputs.configure(yscrollcommand=sb.set)

        r4 = ttk.Frame(self.grp_batch)
        r4.pack(fill=tk.X, **pad)
        ttk.Label(r4, text="Output folder").pack(anchor=tk.W)
        r4b = ttk.Frame(r4)
        r4b.pack(fill=tk.X)
        ttk.Entry(r4b, textvariable=self.var_output_dir).pack(side=tk.LEFT, fill=tk.X, expand=True)
        ttk.Button(r4b, text="Choose…", command=self._choose_outdir).pack(side=tk.LEFT, padx=6)

        # Common options
        grp_opts = ttk.LabelFrame(frm, text="Options")
        grp_opts.pack(fill=tk.X, **pad)
        r5 = ttk.Frame(grp_opts)
        r5.pack(fill=tk.X, **pad)

        ttk.Label(r5, text="Format").grid(row=0, column=0, sticky=tk.W)
        self.cmb_fmt = ttk.Combobox(r5, values=SUPPORTED_FORMATS, textvariable=self.var_format, width=8, state="readonly")
        self.cmb_fmt.grid(row=0, column=1, sticky=tk.W, padx=(8, 16))

        ttk.Label(r5, text="Model").grid(row=0, column=2, sticky=tk.W)
        self.cmb_model = ttk.Combobox(r5, textvariable=self.var_model, values=sorted({DEFAULT_MODEL, "gpt-4o", "gpt-4o-mini"}))
        self.cmb_model.grid(row=0, column=3, sticky="ew", padx=(8, 16))

        ttk.Label(r5, text="PDF DPI").grid(row=0, column=4, sticky=tk.W)
        self.spn_dpi = ttk.Spinbox(r5, from_=100, to=600, increment=10, textvariable=self.var_dpi, width=7)
        self.spn_dpi.grid(row=0, column=5, sticky=tk.W)
        r5.columnconfigure(3, weight=1)

        # API key row
        r6 = ttk.Frame(grp_opts)
        r6.pack(fill=tk.X, **pad)
        ttk.Label(r6, text="OpenAI API Key").grid(row=0, column=0, sticky=tk.W)
        self.ent_key = ttk.Entry(r6, textvariable=self.var_api_key, show="•")
        self.ent_key.grid(row=0, column=1, sticky="ew", padx=(8, 8))
        r6.columnconfigure(1, weight=1)
        self.btn_toggle_key = ttk.Button(r6, text="Show", width=6, command=self._toggle_key)
        self.btn_toggle_key.grid(row=0, column=2, sticky=tk.W)

        # Actions
        rowA = ttk.Frame(frm)
        rowA.pack(fill=tk.X, **pad)
        self.btn_start = ttk.Button(rowA, text="Start OCR", command=self._start_ocr)
        self.btn_start.pack(side=tk.LEFT)
        self.btn_cancel = ttk.Button(rowA, text="Cancel", command=self._cancel, state=tk.DISABLED)
        self.btn_cancel.pack(side=tk.LEFT, padx=(8, 0))
        ttk.Checkbutton(rowA, text="Open when done", variable=self.var_open_when_done).pack(side=tk.RIGHT)

        # Progress + log
        rowB = ttk.Frame(frm)
        rowB.pack(fill=tk.BOTH, expand=True, **pad)
        self.prog = ttk.Progressbar(rowB, mode="indeterminate")
        self.prog.pack(fill=tk.X)
        self.lbl_status = ttk.Label(rowB, text="Ready")
        self.lbl_status.pack(anchor=tk.W, pady=(6, 4))
        self.txt_log = tk.Text(rowB, height=12)
        self.txt_log.pack(fill=tk.BOTH, expand=True)

        # Initial state
        self._toggle_mode()

    # Helpers
    def _toggle_mode(self) -> None:
        """Hide/show single-vs-batch sections without configuring unsupported widget states."""
        batch = self.var_batch.get()
        if batch:
            try:
                self.grp_single.pack_forget()
            except Exception:
                pass
            if not self.grp_batch.winfo_ismapped():
                self.grp_batch.pack(fill=tk.BOTH, expand=True, padx=10, pady=6)
        else:
            try:
                self.grp_batch.pack_forget()
            except Exception:
                pass
            if not self.grp_single.winfo_ismapped():
                self.grp_single.pack(fill=tk.X, padx=10, pady=6)

    def _toggle_key(self) -> None:
        current = self.ent_key.cget("show")
        self.ent_key.configure(show="" if current else "•")
        self.btn_toggle_key.configure(text="Hide" if current else "Show")

    def _choose_input(self) -> None:
        types = [
            ("PDF files", "*.pdf"),
            ("Images", "*.png;*.jpg;*.jpeg;*.bmp;*.tif;*.tiff;*.webp"),
            ("All files", "*.*"),
        ]
        path = filedialog.askopenfilename(title="Select input", filetypes=types)
        if not path:
            return
        self.var_input.set(path)
        if not self.var_output.get().strip():
            self.var_output.set(str(Path(path).with_suffix("." + self.var_format.get())))

    def _choose_output(self) -> None:
        fmt = self.var_format.get()
        initial = (
            self.var_output.get().strip()
            or (self.var_input.get().strip() and str(Path(self.var_input.get()).with_suffix("." + fmt)))
            or "output." + fmt
        )
        path = filedialog.asksaveasfilename(
            title="Save output as",
            defaultextension="." + fmt,
            initialfile=Path(initial).name if initial else None,
            filetypes=[(fmt.upper() + " file", f"*.{fmt}"), ("All Files", "*.*")],
        )
        if path:
            self.var_output.set(path)

    def _choose_outdir(self) -> None:
        d = filedialog.askdirectory(title="Choose output folder")
        if d:
            self.var_output_dir.set(d)

    def _add_files(self) -> None:
        types = [
            ("PDF files", "*.pdf"),
            ("Images", "*.png;*.jpg;*.jpeg;*.bmp;*.tif;*.tiff;*.webp"),
        ]
        paths = filedialog.askopenfilenames(title="Add files", filetypes=types)
        for p in paths:
            self._add_path_if_supported(Path(p))

    def _add_folder(self) -> None:
        d = filedialog.askdirectory(title="Add folder")
        if not d:
            return
        for p in sorted(Path(d).iterdir()):
            if p.is_file():
                self._add_path_if_supported(p)

    def _add_path_if_supported(self, p: Path) -> None:
        if p.suffix.lower() in SUPPORTED_INPUT_EXTS:
            existing = set(self.lst_inputs.get(0, tk.END))
            if str(p) not in existing:
                self.lst_inputs.insert(tk.END, str(p))

    def _remove_selected(self) -> None:
        for i in reversed(self.lst_inputs.curselection()):
            self.lst_inputs.delete(i)

    def _clear_list(self) -> None:
        self.lst_inputs.delete(0, tk.END)

    # Run
    def _start_ocr(self) -> None:
        batch = self.var_batch.get()
        fmt = self.var_format.get()
        model = self.var_model.get().strip() or DEFAULT_MODEL
        dpi = max(72, int(self.var_dpi.get() or DEFAULT_DPI))
        api_key = self.var_api_key.get().strip() or None
        if api_key is None and not os.environ.get("OPENAI_API_KEY"):
            messagebox.showerror("API key required", "Provide an OpenAI API key or set OPENAI_API_KEY.")
            return

        if batch:
            inputs = list(self.lst_inputs.get(0, tk.END))
            if not inputs:
                messagebox.showerror("No files", "Add at least one file to the batch queue.")
                return
            out_dir = self.var_output_dir.get().strip()
            if not out_dir:
                messagebox.showerror("Missing output folder", "Choose an output folder for batch mode.")
                return
            Path(out_dir).mkdir(parents=True, exist_ok=True)
        else:
            inp = self.var_input.get().strip()
            if not inp:
                messagebox.showerror("Missing input", "Choose a PDF or image file.")
                return
            inputs = [inp]
            out = self.var_output.get().strip()
            if not out:
                out = str(Path(inp).with_suffix("." + fmt))
                self.var_output.set(out)
            Path(out).parent.mkdir(parents=True, exist_ok=True)

        # Lock UI
        self._set_running(True)
        self._log("Starting OCR…")
        self.cancel_event.clear()

        def run_job():
            try:
                ocr = OpenAIOCR(model=model, api_key=api_key)
                total_files = len(inputs)

                for file_idx, raw in enumerate(inputs, start=1):
                    if self.cancel_event.is_set():
                        raise RuntimeError("Operation cancelled")
                    inp_path = Path(raw)

                    def on_page(i: int, total_pages: int) -> None:
                        self.log_q.put(f"{inp_path.name}: page {i}/{total_pages}")
                        self.after(0, self._update_page_progress, i, total_pages, file_idx, total_files, inp_path.name)

                    text = ocr_single_file(inp_path, ocr, dpi=dpi, on_page=on_page, cancel=self.cancel_event)

                    if batch:
                        out_path = target_path_for(inp_path, Path(self.var_output_dir.get()), fmt)
                    else:
                        out_path = Path(self.var_output.get())

                    write_output(text, out_path, fmt)
                    self.log_q.put(f"Wrote: {out_path}")

                self.after(0, self._job_done, True, f"Processed {total_files} file(s)")
            except Exception as e:
                self.after(0, self._job_done, False, str(e))

        self.worker = threading.Thread(target=run_job, daemon=True)
        self.worker.start()

    def _update_page_progress(self, i: int, total_pages: int, file_idx: int, total_files: int, name: str) -> None:
        if total_pages <= 0:
            self.prog.configure(mode="indeterminate")
            self.prog.start(10)
        else:
            if self.prog.cget("mode") != "determinate":
                self.prog.configure(mode="determinate")
            self.prog.configure(value=i, maximum=max(1, total_pages))
        self.lbl_status.configure(text=f"{name} — page {i}/{total_pages} (file {file_idx}/{total_files})")

    def _cancel(self) -> None:
        self.cancel_event.set()
        self._log("Cancelling…")

    def _job_done(self, ok: bool, msg: str) -> None:
        self._set_running(False)
        if ok:
            self._log(msg)
            self.lbl_status.configure(text="Done")
            if self.var_open_when_done.get():
                if self.var_batch.get():
                    self._open_path(self.var_output_dir.get().strip())
                else:
                    self._open_path(self.var_output.get().strip())
        else:
            self._log(f"Error: {msg}")
            self.lbl_status.configure(text="Error")
            messagebox.showerror("OCR failed", msg)

    def _open_path(self, path: str) -> None:
        p = Path(path)
        try:
            target = p
            if sys.platform.startswith("win"):
                os.startfile(str(target))  # type: ignore[attr-defined]
            elif sys.platform == "darwin":
                os.system(f"open '{str(target)}'")
            else:
                os.system(f"xdg-open '{str(target)}'")
        except Exception:
            pass

    def _set_running(self, running: bool) -> None:
        state = tk.DISABLED if running else tk.NORMAL
        for w in [self.btn_start, self.cmb_model, self.spn_dpi, self.btn_cancel, self.cmb_fmt]:
            try:
                w.configure(state=(tk.NORMAL if w is self.btn_cancel and running else state))
            except Exception:
                pass
        self.btn_cancel.configure(state=(tk.NORMAL if running else tk.DISABLED))
        if running and self.prog.cget("mode") == "indeterminate":
            self.prog.start(10)
        else:
            self.prog.stop()

    def _log(self, text: str) -> None:
        self.txt_log.insert(tk.END, f"{text}{NL}")
        self.txt_log.see(tk.END)

    def _poll_log(self) -> None:
        try:
            while True:
                line = self.log_q.get_nowait()
                self._log(line)
        except queue.Empty:
            pass
        self.after(100, self._poll_log)


# ---------------------------- CLI ----------------------------

def parse_args(argv: Optional[Sequence[str]] = None) -> argparse.Namespace:
    p = argparse.ArgumentParser(description="OCR images or PDFs to text via OpenAI Vision; double‑click for GUI")
    p.add_argument("inputs", type=Path, nargs="*", help="Input file(s)")
    p.add_argument("-o", "--output", type=Path, default=None, help="Output file (single‑file mode)")
    p.add_argument("--folder", type=Path, default=None, help="Process all supported files in this folder (non‑recursive)")
    p.add_argument("--out-dir", type=Path, default=None, help="Output directory (batch mode)")
    p.add_argument("--format", choices=SUPPORTED_FORMATS, default="txt", help="Output format")
    p.add_argument("--model", default=DEFAULT_MODEL, help=f"OpenAI model (default: {DEFAULT_MODEL})")
    p.add_argument("--dpi", type=int, default=DEFAULT_DPI, help=f"PDF render DPI (default: {DEFAULT_DPI})")
    p.add_argument("--api-key", default=None, help="OpenAI API key (overrides OPENAI_API_KEY)")
    p.add_argument("--gui", action="store_true", help="Force open the GUI")
    p.add_argument("--self-test", action="store_true", help="Run built-in tests and exit")
    return p.parse_args(argv)


def collect_inputs(args: argparse.Namespace) -> List[Path]:
    files: List[Path] = []
    files.extend(args.inputs)
    if args.folder:
        for p in sorted(args.folder.iterdir()):
            if p.is_file() and p.suffix.lower() in SUPPORTED_INPUT_EXTS:
                files.append(p)
    return files


# ---------------------------- Tests ----------------------------

def run_self_tests() -> bool:
    import unittest
    import tempfile

    class OCRUtilsTest(unittest.TestCase):
        def test_page_header(self):
            self.assertEqual(page_header(1), f"{NL}{NL}===== Page 1 ====={NL}")
            self.assertEqual(page_header(12), f"{NL}{NL}===== Page 12 ====={NL}")

        def test_normalize_to_markdown(self):
            raw = "===== Page 1 =====\nHello\n\n===== Page 2 =====\nWorld"
            md = normalize_to_markdown(raw)
            self.assertEqual(md, "## Page 1\nHello\n\n## Page 2\nWorld\n")

        def test_normalize_without_headers(self):
            raw = "Hello\nWorld"
            md = normalize_to_markdown(raw)
            self.assertEqual(md, "Hello\nWorld\n")

        def test_target_path_for(self):
            with tempfile.TemporaryDirectory() as td:
                out = target_path_for(Path("scan.pdf"), Path(td) / "out", "md")
                self.assertEqual(out.name, "scan.md")
                self.assertEqual(out.parent, Path(td) / "out")

        def test_write_output_txt_md(self):
            with tempfile.TemporaryDirectory() as td:
                txtp = Path(td) / "a.txt"
                mdp = Path(td) / "a.md"
                write_output("hello", txtp, "txt")
                self.assertTrue(txtp.exists())
                write_output("===== Page 1 =====\nHello", mdp, "md")
                self.assertTrue(mdp.exists())
                self.assertIn("## Page 1\nHello\n", mdp.read_text())

        def test_newline_constant(self):
            self.assertEqual(NL, "\n")

        @unittest.skipIf(Document is None, "python-docx not installed")
        def test_write_output_docx(self):
            # Smoke test for docx export
            with tempfile.TemporaryDirectory() as td:
                docx_path = Path(td) / "a.docx"
                write_output("===== Page 1 =====\nHello", docx_path, "docx")
                self.assertTrue(docx_path.exists())

    suite = unittest.defaultTestLoader.loadTestsFromTestCase(OCRUtilsTest)
    result = unittest.TextTestRunner(verbosity=2).run(suite)
    return result.wasSuccessful()


# ---------------------------- Entrypoints ----------------------------

def launch_gui() -> None:
    app = OCRApp()
    app.mainloop()


def main(argv: Optional[Sequence[str]] = None) -> int:
    args = parse_args(argv)

    if args.self_test:
        ok = run_self_tests()
        return 0 if ok else 1

    if args.gui or (not args.inputs and not args.folder):
        launch_gui()
        return 0

    files = collect_inputs(args)
    if not files:
        print("[error] No input files provided", file=sys.stderr)
        return 2

    fmt = args.format
    dpi = int(args.dpi)

    try:
        ocr = OpenAIOCR(model=args.model, api_key=args.api_key)
        if len(files) == 1 and not args.folder and args.output:
            text = ocr_single_file(files[0], ocr, dpi=dpi)
            out_path = args.output
            if out_path.suffix.lower() != f".{fmt}":
                out_path = out_path.with_suffix("." + fmt)
            write_output(text, out_path, fmt)
            print(f"[ok] Wrote: {out_path}")
            return 0

        out_dir = args.out_dir or Path.cwd() / "ocr_out"
        out_dir.mkdir(parents=True, exist_ok=True)
        for f in files:
            text = ocr_single_file(f, ocr, dpi=dpi)
            out_path = target_path_for(f, out_dir, fmt)
            write_output(text, out_path, fmt)
            print(f"[ok] Wrote: {out_path}")
        return 0
    except Exception as e:
        print(f"[fatal] {e}", file=sys.stderr)
        return 1


if __name__ == "__main__":  # pragma: no cover
    raise SystemExit(main())

